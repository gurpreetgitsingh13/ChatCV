from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches
from io import BytesIO
import os, json, datetime

app = Flask(__name__)
CORS(app)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(BASE_DIR)
DATA_DIR = os.path.join(ROOT, 'data')
OUTPUT_DIR = os.path.join(ROOT, 'outputs')
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

@app.get("/api/health")
def health():
    return {"ok": True}

@app.post("/api/save_chat_data")
def save_chat_data():
    payload = request.get_json(force=True, silent=True) or {}
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    path = os.path.join(DATA_DIR, f"resume_data_{ts}.json")
    with open(path, "w") as f:
        json.dump(payload, f, indent=2)
    return jsonify({"saved": True, "path": path})

def add_heading(doc, text, size=14, bold=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')
        p_format = p.paragraph_format
        p_format.space_after = Pt(0)

def safe_get(d, k, default=""):
    return (d.get(k) or default)

@app.post("/api/generate_docx")
def generate_docx():
    data = request.get_json(force=True, silent=True) or {}

    doc = Document()

    # Header
    name = safe_get(data, 'full_name', 'Your Name')
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(20)

    contact_bits = list(filter(None, [
        safe_get(data, 'email'), safe_get(data, 'phone'),
        safe_get(data, 'linkedin'), safe_get(data, 'github')
    ]))
    if contact_bits:
        p2 = doc.add_paragraph(" · ".join(contact_bits))
        p2_format = p2.paragraph_format
        p2_format.space_after = Pt(6)

    summary = safe_get(data, 'summary')
    if summary:
        add_heading(doc, "Summary", 14)
        doc.add_paragraph(summary)

    skills = [s.strip() for s in safe_get(data, 'skills', '').split(',') if s.strip()]
    if skills:
        add_heading(doc, "Skills", 14)
        doc.add_paragraph(", ".join(skills))

    # Experience
    work_lines = [x.strip() for x in safe_get(data, 'work','').split('\n') if x.strip()]
    if work_lines:
        add_heading(doc, "Experience", 14)
        for line in work_lines:
            parts = [p.strip() for p in line.split('|')]
            company = parts[0] if len(parts)>0 else ""
            role = parts[1] if len(parts)>1 else ""
            dates = parts[2] if len(parts)>2 else ""
            bullets = parts[3] if len(parts)>3 else ""
            header = f"{role} — {company}"
            if dates: header += f" ({dates})"
            doc.add_paragraph(header)
            bullet_items = [b.strip() for b in bullets.split(';') if b.strip()]
            add_bullets(doc, bullet_items)

    # Projects
    proj_lines = [x.strip() for x in safe_get(data, 'projects','').split('\n') if x.strip()]
    if proj_lines:
        add_heading(doc, "Projects", 14)
        for line in proj_lines:
            parts = [p.strip() for p in line.split('|')]
            title = parts[0] if len(parts)>0 else ""
            stack = parts[1] if len(parts)>1 else ""
            bullets = parts[2] if len(parts)>2 else ""
            header = f"{title} — {stack}"
            doc.add_paragraph(header)
            bullet_items = [b.strip() for b in bullets.split(';') if b.strip()]
            add_bullets(doc, bullet_items)

    # Education
    edu_lines = [x.strip() for x in safe_get(data, 'education','').split('\n') if x.strip()]
    if edu_lines:
        add_heading(doc, "Education", 14)
        for line in edu_lines:
            parts = [p.strip() for p in line.split('|')]
            school = parts[0] if len(parts)>0 else ""
            degree = parts[1] if len(parts)>1 else ""
            dates = parts[2] if len(parts)>2 else ""
            note = parts[3] if len(parts)>3 else ""
            header = f"{school} — {degree}"
            if dates: header += f" ({dates})"
            doc.add_paragraph(header)
            if note:
                doc.add_paragraph(note)

    # Stream result
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return send_file(bio, as_attachment=True, download_name="ChatCV_Resume.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5050"))
    app.run(host="0.0.0.0", port=port)
