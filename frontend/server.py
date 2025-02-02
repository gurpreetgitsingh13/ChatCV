import json
from docx import Document

def create_resume_from_json(json_file, output_file="resume_auto.docx"):
    with open(json_file, "r", encoding="utf-8") as file:
        data = json.load(file)
    
    doc = Document()
    doc.add_heading(data.get("full_name", "Your Name"), level=1)
    doc.add_paragraph(f"Email: {data.get('email', 'your-email@example.com')}")
    doc.add_paragraph(f"Phone: {data.get('phone_number', '123-456-7890')}")
    doc.add_paragraph(f"LinkedIn: {data.get('linkedin', 'linkedin.com/in/yourprofile')}")
    doc.add_paragraph(f"GitHub: {data.get('github', 'github.com/yourprofile')}")
    
    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(data.get("summary", "Your professional summary here."))
    
    doc.add_heading("Skills", level=2)
    skills = data.get("skills", "").split(",")
    for skill in skills:
        doc.add_paragraph(f"• {skill.strip()}")
    
    if data.get("work_experience", "no").lower() == "yes":
        doc.add_heading("Work Experience", level=2)
        experiences = data.get("work_details", "").split(";")
        for experience in experiences:
            doc.add_paragraph(f"• {experience.strip()}")
    
    if data.get("education", "no").lower() == "yes":
        doc.add_heading("Education", level=2)
        education_details = data.get("education_details", "").split(";")
        for education in education_details:
            doc.add_paragraph(f"• {education.strip()}")
    
    if data.get("certifications", "no").lower() == "yes":
        doc.add_heading("Certifications", level=2)
        certifications = data.get("certifications", "").split(",")
        for cert in certifications:
            doc.add_paragraph(f"• {cert.strip()}")
    
    if data.get("projectdetailst", "no").lower() == "yes":
        doc.add_heading("Projects", level=2)
        projects = data.get("project_details", "").split(";")
        for project in projects:
            doc.add_paragraph(f"• {project.strip()}")
    
    doc.save(output_file)
    print(f"✅ Resume saved as {output_file}")

# Example usage
if __name__ == "__main__":
    create_resume_from_json("/Users/geetikadewan/Downloads/resume_data.json")
